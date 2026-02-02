"""
Document parser for extracting text from various file formats
"""
from pathlib import Path
from pypdf import PdfReader
from docx import Document
import re


class DocumentParser:
    """Parse documents to extract text content with location tracking"""

    def __init__(self, filepath):
        """Initialize parser with file path"""
        self.filepath = filepath
        self.full_text = ""
        self.page_breaks = []  # Character positions where pages start
        self.page_count = 0

    def parse_file(self):
        """
        Parse the file and extract text content with location tracking

        Returns:
            Extracted text as string
        """
        ext = Path(self.filepath).suffix.lower()

        if ext == '.pdf':
            return self._parse_pdf()
        elif ext == '.docx':
            return self._parse_docx()
        elif ext == '.doc':
            # Legacy .doc format not supported by python-docx library
            raise ValueError("Legacy .doc files are not supported. Please convert to .docx format using Microsoft Word or LibreOffice.")
        elif ext == '.txt':
            return self._parse_txt()
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _parse_pdf(self):
        """Extract text from PDF with page tracking"""
        try:
            reader = PdfReader(self.filepath)
            pages_text = []
            current_pos = 0

            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text()
                pages_text.append(page_text)

                # Track where this page starts in the full text
                self.page_breaks.append({
                    'page': page_num,
                    'start_pos': current_pos,
                    'end_pos': current_pos + len(page_text)
                })
                current_pos += len(page_text) + 1  # +1 for newline

            self.page_count = len(reader.pages)
            self.full_text = '\n'.join(pages_text)
            return self.full_text
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")

    def _parse_docx(self):
        """Extract text from DOCX with paragraph tracking"""
        try:
            doc = Document(self.filepath)
            paragraphs_text = []
            current_pos = 0

            # DOCX doesn't have page numbers in the same way,
            # so we'll estimate based on paragraph count
            for para_num, paragraph in enumerate(doc.paragraphs, start=1):
                para_text = paragraph.text
                paragraphs_text.append(para_text)

                # Estimate page breaks every ~10 paragraphs (rough approximation)
                if para_num % 10 == 1:
                    estimated_page = (para_num // 10) + 1
                    self.page_breaks.append({
                        'page': estimated_page,
                        'start_pos': current_pos,
                        'end_pos': current_pos + len(para_text)
                    })

                current_pos += len(para_text) + 1  # +1 for newline

            self.page_count = max((len(doc.paragraphs) // 10) + 1, 1)
            self.full_text = '\n'.join(paragraphs_text)
            return self.full_text
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")

    def _parse_txt(self):
        """Extract text from TXT with encoding fallback"""
        # Try multiple encodings in order of likelihood
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(self.filepath, 'r', encoding=encoding) as f:
                    text = f.read()
                    self.full_text = text
                    # TXT files don't have inherent page breaks
                    self.page_count = 1
                    self.page_breaks = [{'page': 1, 'start_pos': 0, 'end_pos': len(text)}]
                    return text
            except UnicodeDecodeError:
                continue
            except Exception as e:
                raise Exception(f"Error parsing TXT: {str(e)}")

        # If all encodings fail, try with error handling
        try:
            with open(self.filepath, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
                self.full_text = text
                self.page_count = 1
                self.page_breaks = [{'page': 1, 'start_pos': 0, 'end_pos': len(text)}]
                return text
        except Exception as e:
            raise Exception(f"Error parsing TXT: {str(e)}")

    def get_page_number(self, char_position):
        """
        Get page number for a given character position

        Args:
            char_position: Character position in the full text

        Returns:
            Page number (1-indexed) or None
        """
        for page_info in self.page_breaks:
            if page_info['start_pos'] <= char_position <= page_info['end_pos']:
                return page_info['page']
        return None

    def find_text_in_document(self, search_text, context_chars=200):
        """
        FEATURE DISABLED: This method is currently not in use.
        Find text in document and return location with context

        Args:
            search_text: Text to search for
            context_chars: Number of characters to show before/after match

        Returns:
            Dictionary with location info, or None if not found
        """
        if not search_text or not self.full_text:
            return None

        # Normalize for searching (case-insensitive)
        text_lower = self.full_text.lower()
        search_lower = search_text.lower().strip()

        # Try exact match first
        index = text_lower.find(search_lower)

        if index == -1:
            # Try fuzzy match - look for substantial overlap
            return self._fuzzy_find_text(search_text, context_chars)

        return self._build_location_info(index, len(search_text), context_chars)

    def _fuzzy_find_text(self, search_text, context_chars=200):
        """
        FEATURE DISABLED: This method is currently not in use.
        Attempt fuzzy matching for text that may be paraphrased

        Args:
            search_text: Text to search for
            context_chars: Number of characters for context

        Returns:
            Location info or None
        """
        # Split search text into words
        words = re.findall(r'\w+', search_text.lower())
        if len(words) < 3:
            return None

        # Look for sequences of at least 3 words
        best_match = None
        best_match_score = 0

        text_lower = self.full_text.lower()

        for i in range(len(words) - 2):
            # Look for 3-word sequences
            three_words = ' '.join(words[i:i+3])
            if three_words in text_lower:
                # Found a match, try to extend it
                start_pos = text_lower.find(three_words)

                # Count how many words match in sequence
                match_score = 3
                search_pos = i + 3
                text_pos = start_pos + len(three_words)

                while search_pos < len(words):
                    next_word = words[search_pos]
                    # Look ahead up to 50 characters for next word
                    next_chunk = text_lower[text_pos:text_pos+50]
                    if next_word in next_chunk:
                        match_score += 1
                        text_pos += next_chunk.find(next_word) + len(next_word)
                        search_pos += 1
                    else:
                        break

                # If this is the best match so far, save it
                if match_score > best_match_score:
                    best_match_score = match_score
                    # Estimate length based on matched words
                    match_length = text_pos - start_pos
                    best_match = self._build_location_info(start_pos, match_length, context_chars)

        # Only return if we matched at least 50% of the words
        if best_match and best_match_score >= len(words) * 0.5:
            best_match['fuzzy_matched'] = True
            best_match['match_confidence'] = min(best_match_score / len(words), 1.0)
            return best_match

        return None

    def _build_location_info(self, start_index, match_length, context_chars):
        """
        FEATURE DISABLED: This method is currently not in use.
        Build location information dictionary
        """
        end_index = start_index + match_length

        # Get page number
        page_num = self.get_page_number(start_index)

        # Extract context
        context_start = max(0, start_index - context_chars)
        context_end = min(len(self.full_text), end_index + context_chars)

        # Get the actual matched text and context
        matched_text = self.full_text[start_index:end_index]
        context_before = self.full_text[context_start:start_index]
        context_after = self.full_text[end_index:context_end]

        # Clean up context (remove excessive whitespace)
        context_before = ' '.join(context_before.split())
        context_after = ' '.join(context_after.split())
        matched_text = ' '.join(matched_text.split())

        return {
            'page': page_num,
            'start_pos': start_index,
            'end_pos': end_index,
            'matched_text': matched_text,
            'context_before': context_before,
            'context_after': context_after,
            'fuzzy_matched': False,
            'match_confidence': 1.0
        }

    @staticmethod
    def parse_file_simple(filepath):
        """
        Simple static method for backward compatibility

        Args:
            filepath: Path to the file

        Returns:
            Extracted text as string
        """
        parser = DocumentParser(filepath)
        return parser.parse_file()
